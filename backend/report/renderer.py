from docx import Document
import os
import subprocess

OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./download_report")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def convert_to_pdf(docx_path):
    try:
        subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", os.path.dirname(docx_path)
        ], check=True)

        return docx_path.replace(".docx", ".pdf")

    except Exception:
        return None
  
import re

def clean_markdown(text: str):
    if not text:
        return ""

    # remove bold/italic
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)

    # convert bullet points
    text = re.sub(r"- ", "• ", text)

    # remove extra hashes
    text = re.sub(r"#+", "", text)

    return text

def render_docx(state):
    filename = f"{state.get('paper_id')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc = Document()

    #title
    doc.add_heading("Research Paper Analysis Report", 0)
    doc.add_paragraph(f"Title: {state.get('title', '')}")

    # SUMMARY
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(state.get("executive_summary", ""))


    # RECOMMENDATION
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(state.get("recommendation", ""))

    # METRICS
    doc.add_heading("Metrics", level=1)
    doc.add_paragraph(f"Consistency: {state.get('consistency_score')}/100")
    doc.add_paragraph(f"Grammar: {state.get('grammar_rating')}")
    doc.add_paragraph(f"Authenticity: {state.get('authenticity_score')}")


    # NOVELTY
    doc.add_heading("Novelty Analysis", level=1)
    doc.add_paragraph(clean_markdown(state.get("novelty")))


    # FACT CHECK
    doc.add_heading("Fact Check", level=1)

    for f in state.get("fact_check", []):
        prefix = "✔ TRUE" if f.get("verified") else "✗ FALSE"
        doc.add_paragraph(f"{prefix} - {f.get('claim')}")

    # SAVE DOCX
    doc.save(filepath)

    return filepath